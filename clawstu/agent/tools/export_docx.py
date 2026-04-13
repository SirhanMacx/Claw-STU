"""Export worksheets and study guides as Word documents.

Uses python-docx. Converts simple markdown to DOCX with basic formatting.
"""

from __future__ import annotations

import re
from pathlib import Path
from typing import Any

from clawstu.agent.base_tool import BaseTool, ToolContext


class ExportDocxTool(BaseTool):
    name = "export_docx"
    description = "Export a markdown file as a Word (.docx) document."
    parameters: dict[str, Any] = {  # noqa: RUF012
        "type": "object",
        "properties": {
            "source_path": {"type": "string", "description": "Path to source .md"},
            "output_path": {"type": "string", "description": "Output .docx path"},
        },
        "required": ["source_path"],
    }

    async def execute(self, args: dict[str, Any], context: ToolContext) -> str:
        from docx import Document  # type: ignore[import-untyped]

        source_path = Path(args["source_path"])
        output_path = Path(
            args.get("output_path") or source_path.with_suffix(".docx"),
        )
        lines = source_path.read_text(encoding="utf-8").splitlines()

        doc = Document()
        for line in lines:
            s = line.strip()
            if s.startswith("# "):
                doc.add_heading(s[2:], level=1)
            elif s.startswith("## "):
                doc.add_heading(s[3:], level=2)
            elif s.startswith("### "):
                doc.add_heading(s[4:], level=3)
            elif re.match(r"^[-*] ", s):
                doc.add_paragraph(s[2:], style="List Bullet")
            elif re.match(r"^\d+\.\s", s):
                doc.add_paragraph(re.sub(r"^\d+\.\s", "", s), style="List Number")
            elif s:
                doc.add_paragraph(s)

        doc.save(str(output_path))
        return f"Exported DOCX: {output_path} ({output_path.stat().st_size} bytes)"
